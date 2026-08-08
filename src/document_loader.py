"""
document_loader.py
-------------------
Extracts raw text content from the four supported course-material file
types (PDF, DOCX, CSV, TXT). Each loader returns plain text; CSV rows are
converted into short natural-language sentences so they embed and retrieve
well semantically (a raw comma-separated row embeds poorly).
"""
from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class RawDocument:
    """A single loaded source file, before chunking."""
    course_code: str
    course_name: str
    file_name: str
    file_type: str          # "pdf" | "docx" | "csv" | "txt"
    text: str
    file_path: str


# ---------------------------------------------------------------------------
# Individual format loaders
# ---------------------------------------------------------------------------

def load_pdf(path: Path) -> str:
    """Extract text from a PDF, page by page."""
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i}]\n{text}")
    return "\n\n".join(pages)


def load_docx(path: Path) -> str:
    """Extract text from a Word document, preserving heading structure as
    plain-text section markers so downstream chunking respects sections."""
    doc = DocxDocument(str(path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = getattr(para.style, "name", None) or ""
        style = style.lower()
        if "heading" in style or "title" in style:
            lines.append(f"\n## {text}\n")
        else:
            lines.append(text)

    # Tables (e.g. grading tables embedded in lecture notes) are flattened
    # into readable rows rather than dropped.
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            lines.append("\n[Table]\n" + "\n".join(rows_text))

    return "\n".join(lines)


def load_csv(path: Path) -> str:
    """Convert a CSV into natural-language sentences (one per row) so each
    row is a self-contained, semantically meaningful chunk candidate."""
    df = pd.read_csv(path)
    df = df.fillna("N/A")
    sentences = []
    for _, row in df.iterrows():
        parts = [f"{col}: {row[col]}" for col in df.columns]
        sentences.append("Record - " + "; ".join(parts) + ".")
    header_note = f"This table has columns: {', '.join(df.columns)}.\n"
    return header_note + "\n".join(sentences)


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".csv": load_csv,
    ".txt": load_txt,
}


# ---------------------------------------------------------------------------
# Course-level loading
# ---------------------------------------------------------------------------

def load_course_materials(course_code: str, course_name: str, folder: Path) -> list[RawDocument]:
    """Load every supported file inside a course folder into RawDocuments."""
    docs: list[RawDocument] = []
    if not folder.exists():
        return docs

    for file_path in sorted(folder.iterdir()):
        suffix = file_path.suffix.lower()
        loader = LOADERS.get(suffix)
        if loader is None:
            continue
        try:
            text = loader(file_path)
        except Exception as exc:  # keep ingestion resilient to one bad file
            text = ""
            print(f"[document_loader] Failed to load {file_path}: {exc}")
        if text.strip():
            docs.append(RawDocument(
                course_code=course_code,
                course_name=course_name,
                file_name=file_path.name,
                file_type=suffix.lstrip("."),
                text=text,
                file_path=str(file_path),
            ))
    return docs


def load_all_courses(course_registry: dict, data_dir: Path) -> list[RawDocument]:
    """Load materials for every course defined in the registry."""
    all_docs: list[RawDocument] = []
    for code, meta in course_registry.items():
        folder = data_dir / meta["folder"]
        all_docs.extend(load_course_materials(code, meta["name"], folder))
    return all_docs
